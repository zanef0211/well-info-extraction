"""
Word文档解析器 - 解析DOCX文件
"""
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    from docx import Document
except ImportError:
    Document = None

from utils.logger import get_logger
from utils.exceptions import ParsingError

logger = get_logger(__name__)


class DocxParser:
    """Word文档解析器 - 使用python-docx解析DOCX文件"""

    def __init__(self):
        """初始化Word解析器"""
        if Document is None:
            raise ParsingError("python-docx未安装,请先安装: pip install python-docx")
        logger.info("Word文档解析器初始化完成")

    def extract_text(self, docx_path: str) -> str:
        """
        提取Word文档中的所有文本

        Args:
            docx_path: Word文档路径

        Returns:
            完整文本内容

        Raises:
            ParsingError: 解析失败时抛出
        """
        try:
            logger.info(f"开始提取Word文档文本: {docx_path}")

            doc = Document(docx_path)

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格文本
            tables_text = self._extract_tables_text(doc)

            # 合并段落和表格文本
            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n".join(tables_text)

            logger.info(f"Word文档文本提取完成,段落数: {len(paragraphs)}, 表格数: {len(doc.tables)}")
            return full_text

        except Exception as e:
            logger.error(f"Word文档文本提取失败: {e}", exc_info=True)
            raise ParsingError(f"Word文档文本提取失败: {str(e)}") from e

    def _extract_tables_text(self, doc: Document) -> List[str]:
        """
        提取表格文本

        Args:
            doc: Document对象

        Returns:
            表格文本列表
        """
        tables_text = []

        for table_idx, table in enumerate(doc.tables):
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))

            tables_text.append(f"\n[表格 {table_idx + 1}]\n" + "\n".join(table_lines))

        return tables_text

    def extract_paragraphs(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        按段落提取文本

        Args:
            docx_path: Word文档路径

        Returns:
            段落信息列表
        """
        try:
            doc = Document(docx_path)
            paragraphs = []

            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "paragraph_number": idx + 1,
                        "text": para.text.strip(),
                        "text_length": len(para.text),
                        "style": para.style.name if para.style else "",
                    })

            logger.info(f"按段落提取完成,共 {len(paragraphs)} 个段落")
            return paragraphs

        except Exception as e:
            logger.error(f"Word文档按段落提取失败: {e}", exc_info=True)
            raise ParsingError(f"Word文档按段落提取失败: {str(e)}") from e

    def extract_tables(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        提取表格数据

        Args:
            docx_path: Word文档路径

        Returns:
            表格数据列表
        """
        try:
            doc = Document(docx_path)
            tables_data = []

            for table_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(cells)

                tables_data.append({
                    "table_number": table_idx + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": rows_data,
                })

            logger.info(f"表格提取完成,共 {len(tables_data)} 个表格")
            return tables_data

        except Exception as e:
            logger.error(f"Word文档表格提取失败: {e}", exc_info=True)
            raise ParsingError(f"Word文档表格提取失败: {str(e)}") from e

    def get_document_info(self, docx_path: str) -> Dict[str, Any]:
        """
        获取Word文档信息

        Args:
            docx_path: Word文档路径

        Returns:
            文档信息字典
        """
        try:
            doc = Document(docx_path)
            path_obj = Path(docx_path)

            return {
                "file_path": docx_path,
                "file_name": path_obj.name,
                "file_size": path_obj.stat().st_size,
                "file_size_mb": round(path_obj.stat().st_size / (1024 * 1024), 2),
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "core_properties": {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "keywords": doc.core_properties.keywords or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                },
            }

        except Exception as e:
            logger.error(f"获取Word文档信息失败: {e}", exc_info=True)
            raise ParsingError(f"获取Word文档信息失败: {str(e)}") from e
